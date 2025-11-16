"""
Document Processor Tool
Process PDFs, Word documents, and images
"""
import os
import base64
from typing import Dict, Any, Optional
from pathlib import Path
from loguru import logger

from core.tool_manager import BaseTool


class DocumentProcessorTool(BaseTool):
    """Document processing tool for various file formats"""
    
    name = "document_processor"
    description = "Process and extract content from PDFs, Word documents, and images"
    
    async def execute(self, file_path: str, operation: str = "read", **kwargs) -> Dict[str, Any]:
        """Execute document processing"""
        operations = {
            "read": self.read_document,
            "extract_text": self.extract_text,
            "get_metadata": self.get_metadata,
            "analyze_image": self.analyze_image
        }
        
        if operation not in operations:
            return {
                "error": f"Unknown operation: {operation}",
                "available": list(operations.keys())
            }
        
        try:
            result = await operations[operation](file_path, **kwargs)
            return {
                "success": True,
                "operation": operation,
                "file": file_path,
                "data": result
            }
        except Exception as e:
            logger.error(f"Document processing error: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
                "file": file_path
            }
    
    async def read_document(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Read and process any supported document type"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self.read_pdf(file_path)
        elif file_ext in ['.docx', '.doc']:
            return await self.read_word(file_path)
        elif file_ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp']:
            return await self.read_image(file_path)
        elif file_ext in ['.txt', '.md', '.py', '.js', '.json', '.csv', '.html', '.css']:
            return await self.read_text_file(file_path)
        else:
            return {"error": f"Unsupported file type: {file_ext}"}
    
    async def read_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        try:
            import PyPDF2
            
            text_content = []
            metadata = {}
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Get metadata
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                    }
                
                # Extract text from all pages
                num_pages = len(pdf_reader.pages)
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text_content.append(f"--- Page {page_num + 1} ---\n{page.extract_text()}")
            
            return {
                "type": "pdf",
                "pages": num_pages,
                "metadata": metadata,
                "content": "\n\n".join(text_content),
                "size_kb": os.path.getsize(file_path) / 1024
            }
        except Exception as e:
            return {"error": f"Failed to read PDF: {str(e)}"}
    
    async def read_word(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # Extract all text
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract tables
            tables_content = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append(table_data)
            
            return {
                "type": "word",
                "paragraphs": len(paragraphs),
                "tables": len(tables_content),
                "content": "\n\n".join(paragraphs),
                "tables_data": tables_content,
                "size_kb": os.path.getsize(file_path) / 1024
            }
        except Exception as e:
            return {"error": f"Failed to read Word document: {str(e)}"}
    
    async def read_image(self, file_path: str) -> Dict[str, Any]:
        """Process image file"""
        try:
            from PIL import Image
            
            img = Image.open(file_path)
            
            # Get image info
            info = {
                "type": "image",
                "format": img.format,
                "mode": img.mode,
                "size": img.size,
                "width": img.width,
                "height": img.height,
                "size_kb": os.path.getsize(file_path) / 1024
            }
            
            # Convert to base64 for display
            with open(file_path, 'rb') as f:
                img_data = base64.b64encode(f.read()).decode('utf-8')
                info["base64"] = f"data:image/{img.format.lower()};base64,{img_data}"
            
            return info
        except Exception as e:
            return {"error": f"Failed to read image: {str(e)}"}
    
    async def read_text_file(self, file_path: str) -> Dict[str, Any]:
        """Read plain text file"""
        try:
            import aiofiles
            import chardet
            
            # Detect encoding
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                detected = chardet.detect(raw_data)
                encoding = detected['encoding'] or 'utf-8'
            
            # Read file
            async with aiofiles.open(file_path, 'r', encoding=encoding) as f:
                content = await f.read()
            
            return {
                "type": "text",
                "encoding": encoding,
                "lines": len(content.split('\n')),
                "content": content,
                "size_kb": os.path.getsize(file_path) / 1024
            }
        except Exception as e:
            return {"error": f"Failed to read text file: {str(e)}"}
    
    async def extract_text(self, file_path: str, **kwargs) -> str:
        """Extract text content from any document"""
        result = await self.read_document(file_path)
        
        if "error" in result:
            return f"Error: {result['error']}"
        
        return result.get("content", str(result))
    
    async def get_metadata(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Get file metadata"""
        try:
            stat = os.stat(file_path)
            
            return {
                "name": os.path.basename(file_path),
                "path": file_path,
                "size_bytes": stat.st_size,
                "size_kb": stat.st_size / 1024,
                "size_mb": stat.st_size / (1024 * 1024),
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "extension": Path(file_path).suffix
            }
        except Exception as e:
            return {"error": str(e)}
    
    async def analyze_image(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """Analyze image properties"""
        try:
            from PIL import Image
            import colorsys
            
            img = Image.open(file_path)
            
            # Get dominant colors
            img_small = img.resize((50, 50))
            colors = img_small.getcolors(2500)
            
            analysis = {
                "dimensions": f"{img.width}x{img.height}",
                "aspect_ratio": round(img.width / img.height, 2),
                "format": img.format,
                "mode": img.mode,
                "is_animated": getattr(img, "is_animated", False),
                "has_transparency": img.mode in ('RGBA', 'LA', 'P')
            }
            
            return analysis
        except Exception as e:
            return {"error": str(e)}


# Standalone function for file upload endpoint
async def process_uploaded_file(file_path: str, file_type: Optional[str] = None) -> Dict[str, Any]:
    """Process uploaded file"""
    tool = DocumentProcessorTool()
    result = await tool.execute(file_path=file_path, operation="read")
    return result

